#!/usr/bin/env python3
"""
DOCX Reader Script for knowing-mcp
Reads Microsoft Word (.docx) files and converts them to text or markdown format
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print(json.dumps({
        "success": False,
        "error": "python-docx library not installed. Run: pip install python-docx"
    }))
    sys.exit(1)


def read_docx_to_text(docx_path):
    """
    Read DOCX file and extract text content
    
    Args:
        docx_path (str): Path to DOCX file
        
    Returns:
        dict: Dictionary containing success status, text content, and metadata
    """
    try:
        docx_path = Path(docx_path)
        
        if not docx_path.exists():
            return {
                "success": False,
                "error": f"DOCX file not found: {docx_path}"
            }
        
        if not docx_path.suffix.lower() in ['.docx', '.doc']:
            return {
                "success": False,
                "error": f"File is not a DOCX/DOC: {docx_path}"
            }
        
        # Read DOCX
        doc = Document(str(docx_path))
        
        # Extract metadata
        core_props = doc.core_properties
        meta_info = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision or 0,
        }
        
        # Extract paragraphs
        paragraphs = []
        paragraph_count = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # Only include non-empty paragraphs
                paragraph_count += 1
                paragraphs.append({
                    "number": paragraph_count,
                    "text": text,
                    "style": para.style.name if para.style else "Normal"
                })
        
        # Extract text from tables
        tables_content = []
        for table_idx, table in enumerate(doc.tables, start=1):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:  # Only include non-empty tables
                tables_content.append({
                    "number": table_idx,
                    "rows": len(table_data),
                    "columns": len(table_data[0]) if table_data else 0,
                    "data": table_data
                })
        
        # Combine all text
        full_text = "\n\n".join([p["text"] for p in paragraphs])
        
        # Add table content to full text
        if tables_content:
            full_text += "\n\n--- Tables ---\n"
            for table in tables_content:
                full_text += f"\n[Table {table['number']}]\n"
                for row in table['data']:
                    full_text += " | ".join(row) + "\n"
        
        return {
            "success": True,
            "file_path": str(docx_path),
            "paragraph_count": paragraph_count,
            "table_count": len(tables_content),
            "metadata": meta_info,
            "text": full_text,
            "paragraphs": paragraphs,
            "tables": tables_content
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": f"Error reading DOCX: {str(e)}"
        }


def read_docx_to_markdown(docx_path):
    """
    Read DOCX file and convert to markdown format
    
    Args:
        docx_path (str): Path to DOCX file
        
    Returns:
        dict: Dictionary containing success status, markdown content, and metadata
    """
    result = read_docx_to_text(docx_path)
    
    if not result.get("success"):
        return result
    
    # Convert to markdown format
    markdown_lines = []
    
    # Add title if available
    if result["metadata"].get("title"):
        markdown_lines.append(f"# {result['metadata']['title']}\n")
    
    # Add metadata section
    markdown_lines.append("## Document Information\n")
    if result["metadata"].get("author"):
        markdown_lines.append(f"**Author:** {result['metadata']['author']}  ")
    if result["metadata"].get("subject"):
        markdown_lines.append(f"**Subject:** {result['metadata']['subject']}  ")
    if result["metadata"].get("keywords"):
        markdown_lines.append(f"**Keywords:** {result['metadata']['keywords']}  ")
    markdown_lines.append(f"**Paragraphs:** {result['paragraph_count']}  ")
    if result["table_count"] > 0:
        markdown_lines.append(f"**Tables:** {result['table_count']}  ")
    if result["metadata"].get("created"):
        markdown_lines.append(f"**Created:** {result['metadata']['created']}  ")
    if result["metadata"].get("modified"):
        markdown_lines.append(f"**Modified:** {result['metadata']['modified']}  ")
    markdown_lines.append("\n---\n")
    
    # Add content by paragraph with style-based formatting
    markdown_lines.append("## Content\n")
    
    current_heading_level = 2
    for para in result["paragraphs"]:
        style = para.get("style", "Normal")
        text = para["text"]
        
        # Convert Word styles to markdown
        if "Heading 1" in style:
            markdown_lines.append(f"\n## {text}\n")
        elif "Heading 2" in style:
            markdown_lines.append(f"\n### {text}\n")
        elif "Heading 3" in style:
            markdown_lines.append(f"\n#### {text}\n")
        elif "Heading 4" in style:
            markdown_lines.append(f"\n##### {text}\n")
        elif "Heading 5" in style or "Heading 6" in style:
            markdown_lines.append(f"\n###### {text}\n")
        elif "Title" in style:
            markdown_lines.append(f"\n# {text}\n")
        elif "Subtitle" in style:
            markdown_lines.append(f"\n*{text}*\n")
        elif "Quote" in style or "Intense Quote" in style:
            markdown_lines.append(f"\n> {text}\n")
        elif "List" in style:
            markdown_lines.append(f"- {text}")
        else:
            markdown_lines.append(f"{text}\n")
    
    # Add tables in markdown format
    if result["tables"]:
        markdown_lines.append("\n---\n")
        markdown_lines.append("## Tables\n")
        
        for table in result["tables"]:
            markdown_lines.append(f"\n### Table {table['number']}\n")
            
            if table["data"]:
                # Add header row
                header = table["data"][0]
                markdown_lines.append("| " + " | ".join(header) + " |")
                markdown_lines.append("|" + "|".join(["---"] * len(header)) + "|")
                
                # Add data rows
                for row in table["data"][1:]:
                    markdown_lines.append("| " + " | ".join(row) + " |")
                
                markdown_lines.append("\n")
    
    markdown_content = "\n".join(markdown_lines)
    
    return {
        "success": True,
        "file_path": result["file_path"],
        "paragraph_count": result["paragraph_count"],
        "table_count": result["table_count"],
        "metadata": result["metadata"],
        "markdown": markdown_content,
        "text": result["text"]
    }


def main():
    parser = argparse.ArgumentParser(description='Read DOCX files and extract text or markdown')
    parser.add_argument('docx_path', help='Path to DOCX file')
    parser.add_argument('--format', choices=['text', 'markdown'], default='markdown',
                        help='Output format (default: markdown)')
    parser.add_argument('--output', '-o', dest='output_path',
                        help='Save output to file instead of printing (useful for large documents)')
    parser.add_argument('--json', action='store_true',
                        help='Output as JSON (for programmatic use)')
    
    args = parser.parse_args()
    
    if args.format == 'markdown':
        result = read_docx_to_markdown(args.docx_path)
    else:
        result = read_docx_to_text(args.docx_path)
    
    if args.json:
        # If output path is specified, save to file and return file info
        if args.output_path and result.get("success"):
            try:
                output_path = Path(args.output_path)
                # Create directory if needed
                output_path.parent.mkdir(parents=True, exist_ok=True)
                
                # Save content to file
                content = result.get("markdown" if args.format == 'markdown' else "text", "")
                output_path.write_text(content, encoding='utf-8')
                
                # Update result to indicate file was saved
                result["saved_to_file"] = str(output_path)
                result["file_size"] = len(content)
                # Remove large content from JSON response
                if "markdown" in result:
                    del result["markdown"]
                if "text" in result:
                    del result["text"]
                
                print(json.dumps(result, indent=2))
            except Exception as e:
                result["success"] = False
                result["error"] = f"Failed to save to file: {str(e)}"
                print(json.dumps(result, indent=2))
                sys.exit(1)
        else:
            print(json.dumps(result, indent=2))
    else:
        if result.get("success"):
            content = result.get("markdown" if args.format == 'markdown' else "text", "")
            
            # Save to file if output path specified
            if args.output_path:
                try:
                    output_path = Path(args.output_path)
                    output_path.parent.mkdir(parents=True, exist_ok=True)
                    output_path.write_text(content, encoding='utf-8')
                    print(f"✅ Saved to: {output_path}")
                except Exception as e:
                    print(f"Error saving to file: {e}", file=sys.stderr)
                    sys.exit(1)
            else:
                print(content)
        else:
            print(f"Error: {result.get('error', 'Unknown error')}", file=sys.stderr)
            sys.exit(1)


if __name__ == "__main__":
    main()
